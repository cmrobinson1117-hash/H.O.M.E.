import streamlit as st
import time
import os
import json
import datetime
from dotenv import load_dotenv

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from openai import OpenAI
import snowflake.connector

# ---------------- LOAD ENV ----------------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ---------------- CONFIG ----------------
st.set_page_config(page_title="H.O.M.E.", layout="wide")

# ---------------- SNOWFLAKE ----------------
def get_connection():
    return snowflake.connector.connect(
        user="YOUR_USER",
        password="YOUR_PASSWORD",
        account="YOUR_ACCOUNT",
        warehouse="YOUR_WAREHOUSE",
        database="YOUR_DB",
        schema="PUBLIC"
    )

# ---------------- AI ANALYSIS ----------------
def analyze_reflection(text):
    prompt = f"""
    Analyze this reflection deeply.

    {text}

    Return JSON:
    {{
      "emotion": "...",
      "theme": "...",
      "psychology": "...",
      "spiritual": "...",
      "action": "..."
    }}
    """

    res = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )

    return res.choices[0].message.content

# ---------------- SAVE ----------------
def save_reflection(user, room, text, video_url=None):
    analysis = analyze_reflection(text)

    timestamp = datetime.datetime.now().isoformat()

    conn = get_connection()
    cur = conn.cursor()

    if video_url:
        try:
            cur.execute("""
                INSERT INTO HOME_REFLECTIONS (USER_ID, ROOM, REFLECTION, INSIGHT, TIMESTAMP, VIDEO_URL)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (user, room, text, analysis, timestamp, video_url))
        except Exception:
            cur.execute("""
                INSERT INTO HOME_REFLECTIONS (USER_ID, ROOM, REFLECTION, INSIGHT, TIMESTAMP)
                VALUES (%s, %s, %s, %s, %s)
            """, (user, room, text, analysis, timestamp))
    else:
        cur.execute("""
            INSERT INTO HOME_REFLECTIONS (USER_ID, ROOM, REFLECTION, INSIGHT, TIMESTAMP)
            VALUES (%s, %s, %s, %s, %s)
        """, (user, room, text, analysis, timestamp))

    cur.close()
    conn.close()

# ---------------- RETRIEVE ----------------
def get_reflections(user):
    conn = get_connection()
    cur = conn.cursor()

    try:
        cur.execute("""
            SELECT ROOM, REFLECTION, TIMESTAMP, INSIGHT, VIDEO_URL
            FROM HOME_REFLECTIONS
            WHERE USER_ID = %s
            ORDER BY TIMESTAMP
        """, (user,))
    except Exception:
        cur.execute("""
            SELECT ROOM, REFLECTION, TIMESTAMP, INSIGHT
            FROM HOME_REFLECTIONS
            WHERE USER_ID = %s
            ORDER BY TIMESTAMP
        """, (user,))

    rows = cur.fetchall()

    cur.close()
    conn.close()

    data = {}
    for row in rows:
        room = row[0]
        if room not in data:
            data[room] = []
        video_url = row[4] if len(row) > 4 else None
        data[room].append({'reflection': row[1], 'timestamp': row[2], 'insight': row[3], 'video_url': video_url})
    return data

# ---------------- PERSONAL PROFILE ----------------
def get_user_profile(user):
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT INSIGHT FROM HOME_REFLECTIONS
        WHERE USER_ID = %s
    """, (user,))

    insights = [row[0] for row in cur.fetchall()]
    cur.close()
    conn.close()

    combined = "\n".join(insights)

    prompt = f"""
    Based on these reflections:

    {combined}

    Identify:
    - dominant emotional pattern
    - recurring struggle
    - growth trend
    - personalized recommendation
    """

    res = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}]
    )

    return res.choices[0].message.content

# ---------------- EMOTION TREND ----------------
def get_emotion_trend(user):
    conn = get_connection()
    cur = conn.cursor()

    cur.execute("""
        SELECT INSIGHT FROM HOME_REFLECTIONS
        WHERE USER_ID = %s
    """, (user,))

    data = [row[0] for row in cur.fetchall()]
    cur.close()
    conn.close()
    return data

# ---------------- HOUSE BUILDER ----------------
def build_house_plan(user, house_name, foundation, blueprint, structure, vision, sustainability):
    reflections = get_reflections(user)
    summary = []
    for room_name, items in reflections.items():
        summary.append(f"{room_name}: {len(items)} reflection(s)")
    reflection_summary = "\n".join(summary) if summary else "No saved reflections yet."

    prompt = f"""
    The user is building a new house plan called '{house_name}'.
    Current house design inputs:
    - Foundation: {foundation}
    - Blueprint: {blueprint}
    - Structure: {structure}
    - Vision: {vision}
    - Sustainability: {sustainability}

    The user has these reflection insights:
    {reflection_summary}

    Create a practical, encouraging house building plan for this user. Include:
    - A clear foundation statement
    - Three design goals
    - A structure checklist
    - A sustainability action step
    - A final inspirational reminder
    """

    try:
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        return res.choices[0].message.content
    except Exception:
        return "Unable to generate a house plan right now. Please try again later."


def show_house_builder(user):
    st.subheader("Build Your Own House")
    st.write("Use your house-building ideas and reflection themes to create a personal home plan.")

    house_name = st.text_input("House Name", value=st.session_state.get("house_name", "My House"), key="house_name")
    foundation = st.text_area("Foundation goals", value=st.session_state.get("house_foundation", ""), key="house_foundation")
    blueprint = st.text_area("Blueprint priorities", value=st.session_state.get("house_blueprint", ""), key="house_blueprint")
    structure = st.text_area("Structure features", value=st.session_state.get("house_structure", ""), key="house_structure")
    vision = st.text_area("Vision & purpose", value=st.session_state.get("house_vision", ""), key="house_vision")
    sustainability = st.text_area("Sustainability actions", value=st.session_state.get("house_sustainability", ""), key="house_sustainability")
    video_url = st.text_input("Video implant URL for this house plan", value=st.session_state.get("house_video_url", ""), key="house_video_url")

    if video_url:
        st.video(video_url)

    if st.button("Generate My House Plan", key="generate_house_plan"):
        plan = build_house_plan(user or "Guest", house_name, foundation, blueprint, structure, vision, sustainability)
        st.markdown("### Your House Plan")
        st.write(plan)
        st.session_state["latest_house_plan"] = plan

    if st.session_state.get("latest_house_plan"):
        st.markdown("### Saved House Plan")
        st.write(st.session_state["latest_house_plan"])

# ---------------- EXPORT ----------------
def export_word(user):
    data = get_reflections(user)
    file = f"{user}_journal.docx"

    doc = Document()
    doc.add_heading("H.O.M.E. Journal", 0)

    for r, reflections in data.items():
        doc.add_heading(r, 1)
        for ref in reflections:
            doc.add_paragraph(f"{ref['timestamp']}: {ref['reflection']}")
            doc.add_paragraph(f"Insight: {ref['insight']}")
            if ref.get('video_url'):
                doc.add_paragraph(f"Video: {ref['video_url']}")

    doc.save(file)
    return file

def export_pdf(user):
    data = get_reflections(user)
    file = f"{user}_journal.pdf"

    doc = SimpleDocTemplate(file)
    styles = getSampleStyleSheet()
    content = []

    for r, reflections in data.items():
        content.append(Paragraph(f"<b>{r}</b>", styles["Heading2"]))
        content.append(Spacer(1, 10))
        for ref in reflections:
            content.append(Paragraph(f"<i>{ref['timestamp']}</i>", styles["BodyText"]))
            content.append(Paragraph(ref['reflection'], styles["BodyText"]))
            content.append(Paragraph(f"<b>Insight:</b> {ref['insight']}", styles["BodyText"]))
            if ref.get('video_url'):
                content.append(Paragraph(f"Video: {ref['video_url']}", styles["BodyText"]))
            content.append(Spacer(1, 10))
        content.append(Spacer(1, 20))

    doc.build(content)
    return file

# ---------------- SESSION ----------------
st.session_state.setdefault("room", "Entry")
st.session_state.setdefault("user", "")

# ---------------- ROOMS CONFIG ----------------
rooms = {
    "Entry": {"next": "Foundation", "prev": None, "verse": "Psalm 84:1-4 - How lovely is your dwelling place, Lord Almighty! My soul yearns, even faints, for the courts of the Lord; my heart and my flesh cry out for the living God. Even the sparrow has found a home, and the swallow a nest for herself, where she may have her young—a place near your altar, Lord Almighty, my King and my God. Blessed are those who dwell in your house; they are ever praising you."},
    "Foundation": {"next": "Blueprint", "prev": "Entry", "verse": "Matthew 7:24-27 - Therefore everyone who hears these words of mine and puts them into practice is like a wise man who built his house on the rock. The rain came down, the streams rose, and the winds blew and beat against that house; yet it did not fall, because it had its foundation on the rock. But everyone who hears these words of mine and does not put them into practice is like a foolish man who built his house on sand. The rain came down, the streams rose, and the winds blew and beat against that house, and it fell with a great crash."},
    "Blueprint": {"next": "Structure", "prev": "Foundation", "verse": "Proverbs 24:3-4 - By wisdom a house is built, and through understanding it is established; through knowledge its rooms are filled with rare and beautiful treasures."},
    "Structure": {"next": "Vision", "prev": "Blueprint", "verse": "Ephesians 2:19-22 - Consequently, you are no longer foreigners and strangers, but fellow citizens with God’s people and also members of his household, built on the foundation of the apostles and prophets, with Christ Jesus himself as the chief cornerstone. In him the whole building is joined together and rises to become a holy temple in the Lord. And in him you also are being built together to become a dwelling in which God lives by his Spirit."},
    "Vision": {"next": "Sustainability", "prev": "Structure", "verse": "Habakkuk 2:2-3 - Then the Lord replied: 'Write down the revelation and make it plain on tablets so that a herald may run with it. For the revelation awaits an appointed time; it speaks of the end and will not prove false. Though it linger, wait for it; it will certainly come and will not delay.'"},
    "Sustainability": {"next": "House Builder", "prev": "Vision", "verse": "Isaiah 40:31 - But those who hope in the Lord will renew their strength. They will soar on wings like eagles; they will run and not grow weary, they will walk and not be faint."},
    "House Builder": {"next": None, "prev": "Sustainability", "verse": "Build a home rooted in purpose, relationships, and action."}
}

# ---------------- UI ----------------
def set_bg(name):
    urls = {
        "Entry": "https://images.unsplash.com/photo-1493809842364-78817add7ffb",
        "Foundation": "https://images.unsplash.com/photo-1505691938895-1758d7feb511",
        "Blueprint": "https://images.unsplash.com/photo-1503387762-592deb58ef4e",
        "Structure": "https://images.unsplash.com/photo-1487958449943-2429e8be8625",
        "Vision": "https://images.unsplash.com/photo-1500530855697-b586d89ba3ee",
        "Sustainability": "https://images.unsplash.com/photo-1507525428034-b723cf961d3e",
        "House Builder": "https://images.unsplash.com/photo-1519710164239-da123dc03ef4"
    }

    st.markdown(f"""
    <style>
    .stApp {{
        background-image: url("{urls.get(name)}");
        background-size: cover;
        animation: fadeIn 1s;
    }}
    @keyframes fadeIn {{from{{opacity:0}} to{{opacity:1}}}}
    </style>
    """, unsafe_allow_html=True)

# ---------------- SIDEBAR ----------------
def sidebar():
    st.sidebar.title("🏠 H.O.M.E.")

    st.sidebar.text_input("Your Name", key="user")

    user = st.session_state.user

    if user:
        if st.sidebar.button("🧠 View My Growth"):
            st.sidebar.write(get_user_profile(user))

        if st.sidebar.button("Export Word"):
            f = export_word(user)
            with open(f, "rb") as file:
                st.sidebar.download_button("Download", file, f)

        if st.sidebar.button("Export PDF"):
            f = export_pdf(user)
            with open(f, "rb") as file:
                st.sidebar.download_button("Download", file, f)

# ---------------- ROOMS ----------------
def entry():
    set_bg("Entry")
    st.title("H.O.M.E.")
    st.write("A house shelters you. A home shapes you.")

    verse = rooms["Entry"]["verse"]
    st.subheader("Welcome Scripture")
    st.write(verse)

    if st.button("Enter"):
        st.session_state.room = "Foundation"

def room(name):
    set_bg(name)
    st.title(name)

    verse = rooms[name]["verse"]
    st.subheader("Scripture for Reflection")
    st.write(verse)

    if name == "House Builder":
        show_house_builder(st.session_state.user)
        return

    if st.session_state.user:
        data = get_reflections(st.session_state.user)
        if name in data:
            st.subheader("Previous Reflections")
            for ref in data[name]:
                with st.expander(f"Reflection on {ref['timestamp']}"):
                    st.write(ref['reflection'])
                    st.write(f"**Insight:** {ref['insight']}")
                    if ref.get('video_url'):
                        st.video(ref['video_url'])
    else:
        st.warning("Please enter your name in the sidebar to view and save reflections.")

    text = st.text_area("Add your reflection", key=f"reflection_{name}")
    video_url = st.text_input("Video implant URL for this reflection", key=f"video_{name}")

    if st.button("Save Reflection", key=f"save_{name}") and text and st.session_state.user:
        save_reflection(st.session_state.user, name, text, video_url or None)
        st.success("Reflection saved!")
        st.experimental_rerun()

    if st.button("✨ Generate Insight", key=f"insight_{name}") and text:
        insight = analyze_reflection(text)
        st.write(insight)

    col1, col2 = st.columns(2)
    with col1:
        if rooms[name]["prev"] and st.button("Previous"):
            st.session_state.room = rooms[name]["prev"]
    with col2:
        if rooms[name]["next"] and st.button("Next"):
            st.session_state.room = rooms[name]["next"]

# ---------------- RUN ----------------
sidebar()

if st.session_state.room == "Entry":
    entry()
elif st.session_state.room == "Foundation":
    room("Foundation")
elif st.session_state.room == "Blueprint":
    room("Blueprint")
elif st.session_state.room == "Structure":
    room("Structure")
elif st.session_state.room == "Vision":
    room("Vision")
elif st.session_state.room == "Sustainability":
    room("Sustainability")
elif st.session_state.room == "House Builder":
    room("House Builder")
